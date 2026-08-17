from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from .config import settings
from .cleaner import clean_text, detect_and_decode
from .cleaner import deduplicate_chunks
from .models import Chunk, DocumentRecord
from .storage import Database, _new_id, _now
from .text import chunk_text


def extract_text(path: Path, filename: str = "") -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
        return clean_text(detect_and_decode(path.read_bytes()))
    if suffix in {".docx"}:
        from docx import Document

        doc = Document(path)
        parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return clean_text("\n\n".join(parts))
    if suffix in {".pdf"}:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return clean_text("\n\n".join(pages))
    raise ValueError(f"暂不支持的文件类型：{suffix or filename}")


def normalize_tags(tags: list[str] | str | None) -> list[str]:
    if tags is None:
        return []
    if isinstance(tags, str):
        tags = re.split(r"[,\u3001;；\s]+", tags)
    return sorted({tag.strip() for tag in tags if tag and tag.strip()})


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def build_document_record(kb_id: str, file_path: Path, text: str) -> DocumentRecord:
    return _build_document_record(kb_id, file_path, text, "public")


def build_document_record_with_mode(kb_id: str, file_path: Path, text: str, access_mode: str) -> DocumentRecord:
    return _build_document_record(kb_id, file_path, text, access_mode)


def _build_document_record(kb_id: str, file_path: Path, text: str, access_mode: str) -> DocumentRecord:
    return DocumentRecord(
        id=_new_id("doc"),
        kb_id=kb_id,
        name=file_path.name,
        file_path=str(file_path),
        content_hash=hash_text(text),
        status="ready",
        access_mode=access_mode,
        created_at=_now(),
    )


def build_chunks(kb_id: str, document: DocumentRecord, text: str, tags: list[str]) -> list[Chunk]:
    parts = chunk_text(text, chunk_size=settings.chunk_size, overlap=settings.chunk_overlap)
    parts = deduplicate_chunks(parts)
    chunks: list[Chunk] = []
    for index, part in enumerate(parts):
        chunks.append(
            Chunk(
                id=_new_id("chunk"),
                kb_id=kb_id,
                document_id=document.id,
                document_name=document.name,
                text=part,
                index=index,
                tags=tags,
            )
        )
    return chunks


def ingest_file(db: Database, kb_id: str, file_path: Path, tags: list[str] | None = None) -> tuple[DocumentRecord, list[Chunk]]:
    return ingest_file_with_mode(db, kb_id, file_path, tags, "public")


def ingest_file_with_mode(
    db: Database,
    kb_id: str,
    file_path: Path,
    tags: list[str] | None = None,
    access_mode: str = "public",
) -> tuple[DocumentRecord, list[Chunk]]:
    text = extract_text(file_path)
    record = build_document_record_with_mode(kb_id, file_path, text, access_mode)
    existing = db.get_document_by_hash(kb_id, record.content_hash)
    if existing:
        raise FileExistsError(f"文档已存在：{existing.name}")

    tags = normalize_tags(tags)
    chunks = build_chunks(kb_id, record, text, tags)
    db.add_document(record)
    db.replace_chunks(record.id, chunks)
    return record, chunks


def delete_document(db: Database, vector_store: Any, kb_id: str, document_id: str) -> None:
    db.delete_document(document_id)
    vector_store.delete_document(kb_id, document_id)

